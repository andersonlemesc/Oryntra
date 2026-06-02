"""Text extraction: direct decode for text/markdown/csv, pypdf for digital PDFs,
and a vision-LLM fallback for scanned/complex PDFs.

The vision fallback rasterizes pages with pypdfium2 (CPU, no GPU) and sends them
to a BYOK multimodal model. Heavy/blocking work runs in a thread, and page count
is capped so a single document can never exhaust the container.
"""

from __future__ import annotations

import asyncio
import base64
import io
import logging
from dataclasses import dataclass
from typing import Any

logger = logging.getLogger(__name__)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PLAIN_TEXT_MIMES: frozenset[str] = frozenset(
    {
        "text/plain",
        "text/markdown",
        "text/x-markdown",
        "text/csv",
        "application/json",
    }
)

VISION_PROVIDERS: frozenset[str] = frozenset({"openai", "anthropic", "gemini"})

MIN_PDF_TEXT_CHARS = 100
VISION_MAX_PAGES = 50
RENDER_SCALE = 2.0

# Caps to bound work / memory on untrusted uploads.
PDF_TEXT_MAX_PAGES = 500
# Reject decompression bombs: a small docx/zip can inflate to gigabytes of XML.
DOCX_MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024

_VISION_PROMPT = (
    "Transcreva todo o conteúdo textual desta página de documento em Markdown. "
    "Preserve títulos, listas e tabelas. Não adicione comentários nem explicações; "
    "responda apenas com o conteúdo transcrito."
)


class UnsupportedMimeTypeError(ValueError):
    pass


class VisionUnavailableError(RuntimeError):
    pass


@dataclass
class VisionCredential:
    provider: str
    model: str
    api_key: str
    base_url: str | None = None


@dataclass
class ExtractResult:
    text: str
    method: str  # "text" | "pdf_lib" | "pdf_vision" | "docx"


def _decode(data: bytes) -> str:
    return data.decode("utf-8", errors="replace").strip()


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = [page.extract_text() or "" for page in reader.pages[:PDF_TEXT_MAX_PAGES]]
    return "\n\n".join(parts).strip()


def _docx_text(data: bytes) -> str:
    import zipfile

    import docx

    # docx is a zip; guard against decompression bombs before parsing.
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            total = sum(info.file_size for info in archive.infolist())
    except zipfile.BadZipFile as exc:
        raise UnsupportedMimeTypeError("invalid docx (not a zip archive)") from exc

    if total > DOCX_MAX_UNCOMPRESSED_BYTES:
        raise VisionUnavailableError("docx uncompressed size exceeds the allowed limit")

    document = docx.Document(io.BytesIO(data))
    return "\n\n".join(paragraph.text for paragraph in document.paragraphs).strip()


def _render_pdf_pages(data: bytes, max_pages: int) -> list[bytes]:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(data)
    try:
        images: list[bytes] = []
        for index in range(min(len(pdf), max_pages)):
            page = pdf[index]
            bitmap = page.render(scale=RENDER_SCALE)
            pil_image = bitmap.to_pil()
            buffer = io.BytesIO()
            pil_image.save(buffer, format="PNG")
            images.append(buffer.getvalue())
        return images
    finally:
        pdf.close()


def _build_vision_model(cred: VisionCredential) -> Any:
    base_url = cred.base_url or None

    if cred.provider == "openai":
        from langchain_openai import ChatOpenAI

        kwargs: dict[str, Any] = {}
        if base_url is not None:
            kwargs["base_url"] = base_url
        return ChatOpenAI(model=cred.model, api_key=cred.api_key, temperature=0, **kwargs)

    if cred.provider == "anthropic":
        from langchain_anthropic import ChatAnthropic

        kwargs = {}
        if base_url is not None:
            kwargs["base_url"] = base_url
        return ChatAnthropic(
            model_name=cred.model,
            api_key=cred.api_key,
            temperature=0,
            timeout=None,
            stop=None,
            **kwargs,
        )

    if cred.provider == "gemini":
        from langchain_google_genai import ChatGoogleGenerativeAI

        kwargs = {}
        if base_url is not None:
            kwargs["client_options"] = {"api_endpoint": base_url}
        return ChatGoogleGenerativeAI(
            model=cred.model,
            google_api_key=cred.api_key,
            temperature=0,
            **kwargs,
        )

    raise VisionUnavailableError(f"provider '{cred.provider}' does not support vision")


async def _pdf_vision(data: bytes, cred: VisionCredential) -> str:
    from langchain_core.messages import HumanMessage

    images = await asyncio.to_thread(_render_pdf_pages, data, VISION_MAX_PAGES)
    if not images:
        return ""

    chat_model = _build_vision_model(cred)
    pages: list[str] = []

    for image in images:
        data_url = "data:image/png;base64," + base64.b64encode(image).decode("ascii")
        message = HumanMessage(
            content=[
                {"type": "text", "text": _VISION_PROMPT},
                {"type": "image_url", "image_url": {"url": data_url}},
            ]
        )
        response = await chat_model.ainvoke([message])
        content = response.content if isinstance(response.content, str) else ""
        content = (content or "").strip()
        if content:
            pages.append(content)

    return "\n\n".join(pages)


async def extract_text(
    *,
    data: bytes,
    mime: str,
    vision: VisionCredential | None = None,
) -> ExtractResult:
    normalized = (mime or "").split(";")[0].strip().lower()

    if normalized in PLAIN_TEXT_MIMES or normalized.startswith("text/"):
        return ExtractResult(_decode(data), "text")

    if normalized == PDF_MIME:
        text = _pdf_text(data)
        if len(text) >= MIN_PDF_TEXT_CHARS:
            return ExtractResult(text, "pdf_lib")

        if vision is not None and vision.provider in VISION_PROVIDERS:
            vision_text = await _pdf_vision(data, vision)
            if vision_text:
                return ExtractResult(vision_text, "pdf_vision")

        if text:
            return ExtractResult(text, "pdf_lib")

        raise VisionUnavailableError(
            "PDF has little to no extractable text and no usable vision extractor was provided"
        )

    if normalized == DOCX_MIME:
        return ExtractResult(_docx_text(data), "docx")

    raise UnsupportedMimeTypeError(f"unsupported mime type: {normalized}")
